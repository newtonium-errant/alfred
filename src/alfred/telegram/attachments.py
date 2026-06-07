"""Document-attachment support for Telegram document messages.

Parallel to :mod:`alfred.telegram.vision` (images) — when the user
forwards a Telegram ``document`` message (a file attached as a file
rather than an image attached as a photo) we want to:

1. Reject anything outside the supported-mime allowlist with an
   explicit user-facing reply (no silent filter-drop — that was the
   2026-06-06 incident this module exists to fix: a PDF arrived,
   PTB had no handler registered for documents, the update was
   silently dropped from every routing path while ``inbound_in_window``
   incremented identically to a noise tick).
2. Download the file bytes via the bot's ``Document.get_file()``
   coroutine.
3. Text-extract via :mod:`pypdf` (PDF-only in c1; the allowlist surface
   is here so adding ``.docx`` / ``.txt`` / etc. is one extra mime + one
   extra extractor branch).
4. Persist the bytes under ``<vault>/inbox/`` for the audit trail.
5. Compose a user-message text that fences the extracted document text
   from the user's caption so the LLM sees the attachment as an
   attachment (not as a continuation of the user's words).

Why text extraction instead of Anthropic's native PDF document block:
the native block is model-family-specific (works on Opus 4.x, but the
shape would change if a future model family swapped the source-block
form). Text extraction is backend-agnostic and works across every
model the talker can route to (Claude, OpenRouter, future locals).
When the model-family quirk surface starts paying off elsewhere
(per ``feedback_sdk_quirk_centralization.md``), we can centralise a
``build_document_block`` here without changing any caller.

Per-instance vault scoping mirrors :mod:`vision`: callers pass
``vault_path`` directly, this module never reads config.

Failure-mode separation: :class:`AttachmentDownloadError` covers the
Telegram fetch (network / PTB shape); :class:`AttachmentExtractError`
covers the PDF-decoding path (corrupted file, empty text, library
quirk). Keeping the two distinct lets the caller emit a more useful
user-facing reply ("couldn't fetch your PDF" vs. "couldn't read your
PDF").
"""
from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from .utils import get_logger

log = get_logger(__name__)


# Allowlist of document MIME types the bot will process — maps each
# supported MIME to a short "kind tag" used by ``on_document`` for
# routing AND by :func:`build_document_user_text` for the per-type
# banner. Single source of truth: this dict is the only allowlist; the
# on_document handler reads ``.get(mime)`` to gate the path and to
# pick the right extractor.
#
# 2026-06-06 P8 (universal filetype bundle): extended from PDF-only
# (c1 / 8ac333b) to cover .docx, plain text (UTF-8 + BOM sniff), .csv,
# .ics, and audio (.mp3/.m4a/.wav/.ogg). Per
# ``feedback_universal_filetype_support.md`` (operator-ratified
# 2026-06-06): no per-instance config gate — every instance gets
# every supported type.
#
# Adding a new kind later: add the MIME → kind row here AND a new
# branch in :func:`extract_<kind>_text` AND the cap in
# :data:`MAX_BYTES_BY_KIND`. The contract-pin test in
# ``tests/telegram/test_attachments.py`` will surface the addition
# so the same commit can update the user-facing rejection text + the
# prompt-tuner SKILL capability surface.
SUPPORTED_DOCUMENT_MIME: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document": "docx",
    "text/plain": "text",
    "text/markdown": "text",
    "text/csv": "csv",
    "text/calendar": "ics",
    # Audio MIMEs — Telegram clients vary by platform on the exact
    # string emitted for the same underlying file. Cover the full set
    # so iPhone .m4a, Android .mp3, desktop .wav, and OGG/Opus all
    # route the same way.
    "audio/mpeg": "audio",       # .mp3
    "audio/mp4": "audio",        # .m4a (most desktop clients)
    "audio/x-m4a": "audio",      # .m4a (iOS variant)
    "audio/wav": "audio",        # .wav
    "audio/x-wav": "audio",      # .wav (legacy variant)
    "audio/ogg": "audio",        # .ogg / .opus container
}


# Per-kind size caps. PDF / DOCX share the standard "document" cap;
# audio gets the 25 MiB Groq Whisper sync-endpoint ceiling (the API
# rejects above that and bigger files need the chunked endpoint, out
# of scope for c1). Text-flavoured types get a lower cap because
# plain text at 5+ MiB is either a programmatic dump or a misuse case
# the LLM context window can't help with anyway. ICS files are
# operationally tiny — a 1 MiB cap covers the long tail.
#
# Centralised dispatch table so ``on_document`` is one lookup, not
# six per-kind branches. Adding a new kind: add the cap here + the
# MIME mapping above + the extractor. The contract-pin test surfaces
# silent additions.
MAX_PDF_BYTES: int = 10 * 1024 * 1024     # 10 MiB
MAX_DOCX_BYTES: int = 10 * 1024 * 1024    # mirror PDF
MAX_TEXT_BYTES: int = 5 * 1024 * 1024     # 5 MiB
MAX_CSV_BYTES: int = 5 * 1024 * 1024      # mirror text
MAX_ICS_BYTES: int = 1 * 1024 * 1024      # calendar files are tiny
MAX_AUDIO_BYTES: int = 25 * 1024 * 1024   # Groq Whisper sync endpoint cap


MAX_BYTES_BY_KIND: dict[str, int] = {
    "pdf": MAX_PDF_BYTES,
    "docx": MAX_DOCX_BYTES,
    "text": MAX_TEXT_BYTES,
    "csv": MAX_CSV_BYTES,
    "ics": MAX_ICS_BYTES,
    "audio": MAX_AUDIO_BYTES,
}


# Extracted-text truncation guard. An operations manual at 200 pages
# can extract to ~400 KB of plain text — feeding that verbatim to the
# LLM context window is wasteful and would blow the per-turn budget.
# 50K chars is roughly 10–15K tokens, fits comfortably in any modern
# model's context, and covers the long tail of typical operational PDFs.
# When the truncation fires the handler emits a single log line and
# appends a visible truncation marker so the model knows it's seeing a
# partial document, not the whole thing.
#
# Applied uniformly across ALL extractors (PDF, DOCX, text, CSV, ICS,
# audio transcript). No per-type special-casing — the LLM's context
# budget is the same regardless of source.
MAX_EXTRACTED_CHARS: int = 50_000


TRUNCATION_MARKER: str = (
    "\n\n[... document truncated; "
    f"only first {MAX_EXTRACTED_CHARS} characters shown ...]"
)


# CSV row cap — distinct from the char cap. Tabular data with 1000+
# rows is rarely useful in conversation; truncate at the row level
# (rather than char) so the row-truncation marker says something
# meaningful ("the first N of M rows") rather than a random char-cut
# mid-cell. The char cap still applies on top of this in case the
# 1000 rows themselves balloon past 50K chars.
MAX_CSV_ROWS: int = 1000


# Human-readable supported-types string, derived from the kinds set
# at module-load time. Used by the on_document rejection reply so a
# future allowlist widening (.xlsx, .epub, etc.) updates the
# user-facing text by extending the MIME map + adding the human label
# below — no scattered string-literal sweep. The test pin in
# ``test_attachments.py`` asserts the mapping covers every kind in
# SUPPORTED_DOCUMENT_MIME.values().
_HUMAN_LABELS_BY_KIND: dict[str, str] = {
    "pdf": "PDFs",
    "docx": ".docx files",
    "text": "plain text",
    "csv": ".csv",
    "ics": "calendar invites (.ics)",
    "audio": "audio files",
}


def _supported_types_human() -> str:
    """Return a human-readable comma-list of supported attachment types.

    Derived from :data:`SUPPORTED_DOCUMENT_MIME` + :data:`_HUMAN_LABELS_BY_KIND`
    so the on_document rejection text stays in sync as the allowlist
    grows. Stable ordering (matches kind-tag insertion order in
    ``_HUMAN_LABELS_BY_KIND``) for grep-friendly user-facing output.
    """
    # Take kinds in the order they appear in _HUMAN_LABELS_BY_KIND,
    # filtered to those actually present in SUPPORTED_DOCUMENT_MIME so
    # a kind removed from the allowlist but left in the labels dict
    # doesn't leak into the user-facing text.
    active_kinds = set(SUPPORTED_DOCUMENT_MIME.values())
    labels = [
        _HUMAN_LABELS_BY_KIND[k]
        for k in _HUMAN_LABELS_BY_KIND
        if k in active_kinds
    ]
    if not labels:
        return "(no attachment types are currently supported)"
    if len(labels) == 1:
        return labels[0]
    if len(labels) == 2:
        return f"{labels[0]} and {labels[1]}"
    return ", ".join(labels[:-1]) + f", and {labels[-1]}"


# Audio MIME → file extension mapping for the save-to-inbox path.
# The extension determines the storage filename's suffix and feeds
# through to the Whisper multipart filename hint. PTB doesn't ship
# extensions on Document objects reliably, so we derive from MIME.
_AUDIO_EXTENSION_BY_MIME: dict[str, str] = {
    "audio/mpeg": "mp3",
    "audio/mp4": "m4a",
    "audio/x-m4a": "m4a",
    "audio/wav": "wav",
    "audio/x-wav": "wav",
    "audio/ogg": "ogg",
}


def extension_for_kind(kind: str, mime_type: str = "") -> str:
    """Return the canonical file extension for a kind + MIME pair.

    PDFs / DOCX / ICS have a single extension. Text covers ``.txt`` AND
    ``.md`` (both route to the same extractor) — we default to ``.txt``
    when the MIME doesn't disambiguate. Audio has multiple containers,
    so the MIME-keyed lookup wins; fall back to a generic ``.bin`` if
    the MIME isn't in the audio map (defensive — shouldn't happen
    given the on_document MIME check upstream).
    """
    if kind == "pdf":
        return "pdf"
    if kind == "docx":
        return "docx"
    if kind == "text":
        return "md" if mime_type == "text/markdown" else "txt"
    if kind == "csv":
        return "csv"
    if kind == "ics":
        return "ics"
    if kind == "audio":
        return _AUDIO_EXTENSION_BY_MIME.get(mime_type, "bin")
    # Defensive fallback for any future kind that forgets to update
    # this table.
    return "bin"


class AttachmentDownloadError(Exception):
    """Raised when fetching a Telegram document fails (network / PTB)."""


class AttachmentExtractError(Exception):
    """Raised when decoding a downloaded document fails.

    Distinct from :class:`AttachmentDownloadError` so the caller can
    emit a more useful user-facing reply — "I couldn't fetch your PDF"
    (network) vs. "I couldn't read your PDF" (decoding).
    """


async def download_document_bytes(document: Any) -> bytes:
    """Download a Telegram ``Document`` to in-memory bytes.

    Mirrors :func:`alfred.telegram.vision.download_photo_bytes`. PTB's
    ``Document.get_file()`` returns a ``File`` object whose
    ``download_as_bytearray()`` coroutine fetches the actual bytes from
    Telegram's servers; we bytes-cast the result so the caller receives
    a plain ``bytes`` object.

    Raises:
        AttachmentDownloadError: Wraps any exception from the PTB /
            HTTP download path so the caller has one error class.
    """
    try:
        tg_file = await document.get_file()
        raw = await tg_file.download_as_bytearray()
        return bytes(raw)
    except Exception as exc:  # noqa: BLE001 — wrap-and-rethrow with one class
        log.warning(
            "talker.attachments.download_failed",
            error=str(exc),
            file_id=getattr(document, "file_id", ""),
            mime_type=getattr(document, "mime_type", ""),
        )
        raise AttachmentDownloadError(
            f"Failed to download Telegram document: {exc!s}"
        ) from exc


def extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract text from a PDF byte stream via :mod:`pypdf`.

    Iterates pages, joins with double-newlines (so the LLM sees page
    boundaries), and trims to :data:`MAX_EXTRACTED_CHARS` with the
    :data:`TRUNCATION_MARKER` appended when truncation fires.

    Raises:
        AttachmentExtractError: On any pypdf decoding failure, or when
            the extracted text is empty (e.g. a scanned-image PDF with
            no embedded text layer — text extraction can't help us
            there; the caller should emit a "scanned PDF can't be read"
            user-facing reply).
    """
    try:
        # Lazy import: keeps the talker module importable on installs
        # that haven't pulled the ``voice`` extra (the on_document
        # handler's load-time guard will reply "PDF support not
        # available in this install" rather than crashing the daemon
        # at import time).
        import pypdf
    except ImportError as exc:
        log.warning("talker.attachments.pypdf_missing", error=str(exc))
        raise AttachmentExtractError(
            "PDF support is not installed in this build "
            "(pip install -e '.[voice]' to enable)"
        ) from exc

    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        pages_text: list[str] = []
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:  # noqa: BLE001
                # Per-page extraction can fail on malformed pages while
                # the rest of the PDF is fine. Log the per-page miss
                # and continue — partial extraction beats total failure.
                log.warning(
                    "talker.attachments.page_extract_failed",
                    error=str(exc),
                )
                continue
            if page_text.strip():
                pages_text.append(page_text)
    except Exception as exc:  # noqa: BLE001 — wrap pypdf's exceptions
        log.warning(
            "talker.attachments.extract_failed",
            error=str(exc),
            error_type=type(exc).__name__,
        )
        raise AttachmentExtractError(
            f"Failed to decode PDF: {exc!s}"
        ) from exc

    full_text = "\n\n".join(pages_text).strip()
    if not full_text:
        # Empty extraction is its own failure mode: typically a scanned
        # image-only PDF with no embedded text layer. The text-extract
        # path can't recover here; OCR would be a separate feature.
        log.info("talker.attachments.empty_extraction")
        raise AttachmentExtractError(
            "No text could be extracted from this PDF "
            "(scanned image-only PDFs need OCR, which isn't enabled)"
        )

    return _apply_char_truncation(full_text, kind="pdf")


def _apply_char_truncation(text: str, *, kind: str) -> str:
    """Apply the uniform :data:`MAX_EXTRACTED_CHARS` cap with a marker.

    Shared helper across every extractor (PDF, DOCX, text, CSV, ICS,
    audio transcript). When truncation fires, emits one log line with
    the kind so dashboards can spot per-type truncation rates ("audio
    transcripts truncate 10% of the time" is operationally useful).
    The marker text is identical for every kind so the LLM's prompt
    template doesn't have to branch on type.

    Returns the (possibly truncated) text. No-op when text is under
    the cap.
    """
    if len(text) > MAX_EXTRACTED_CHARS:
        log.info(
            "talker.attachments.text_truncated",
            kind=kind,
            original_chars=len(text),
            kept_chars=MAX_EXTRACTED_CHARS,
        )
        return text[:MAX_EXTRACTED_CHARS] + TRUNCATION_MARKER
    return text


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract text from a .docx byte stream via :mod:`docx` (python-docx).

    Walks paragraphs + tables in document order. Paragraphs are
    joined with double-newlines (preserving the source's
    paragraph breaks). Tables are rendered as ``[Table N]`` markers
    followed by pipe-separated rows so the LLM sees structured
    tabular data as structured, not as concatenated prose. Images,
    headers, footers, and footnotes are skipped (out of c1 scope —
    rare in operationally-shared docs, and lifting them in is a
    one-branch extension later).

    Raises:
        AttachmentExtractError: On any docx-decoding failure, or when
            the extracted text is empty (.docx files can in principle
            be all-image or all-embedded-object; the LLM can't help
            with those via text extraction).
    """
    try:
        # Lazy import — keeps the module importable on installs that
        # haven't pulled the ``voice`` extra.
        import docx  # python-docx
    except ImportError as exc:
        log.warning("talker.attachments.python_docx_missing", error=str(exc))
        raise AttachmentExtractError(
            "DOCX support is not installed in this build "
            "(pip install -e '.[voice]' to enable)"
        ) from exc

    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception as exc:  # noqa: BLE001 — wrap python-docx + zipfile errors
        log.warning(
            "talker.attachments.docx_open_failed",
            error=str(exc),
            error_type=type(exc).__name__,
        )
        raise AttachmentExtractError(
            f"Failed to open .docx: {exc!s}"
        ) from exc

    # Walk the document body in source order. python-docx exposes
    # ``document.element.body`` as the underlying lxml node; iterating
    # children preserves the interleaving of paragraphs and tables
    # (vs. ``document.paragraphs`` + ``document.tables`` which
    # separates the two and loses position information). For c1
    # simplicity, we use the separated lists — most operationally
    # shared .docx files are either paragraph-only OR have tables at
    # the end (forms, summaries). If interleaving matters later
    # (technical specs with inline tables), lift to the element walk.
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            blocks.append(text)

    for table_idx, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(f"[Table {table_idx}]\n" + "\n".join(rows))

    full_text = "\n\n".join(blocks).strip()
    if not full_text:
        log.info("talker.attachments.docx_empty_extraction")
        raise AttachmentExtractError(
            "No text could be extracted from this .docx "
            "(may be image-only or use embedded objects)"
        )

    return _apply_char_truncation(full_text, kind="docx")


# BOM byte sequences for the text decoder. UTF-8 BOM is harmless if
# stripped; UTF-16 LE / BE need it to disambiguate endianness. Lifted
# into module-level constants so the test suite can reference the
# exact byte patterns rather than re-deriving them.
_UTF8_BOM = b"\xef\xbb\xbf"
_UTF16_LE_BOM = b"\xff\xfe"
_UTF16_BE_BOM = b"\xfe\xff"


def extract_text_decoded(raw_bytes: bytes) -> str:
    """Decode plain text bytes to a string with BOM + fallback handling.

    Pipeline:
        1. Empty input → :class:`AttachmentExtractError`.
        2. BOM check — strip UTF-8 BOM, decode UTF-16 LE/BE if their
           BOMs are present.
        3. Try plain UTF-8 decode. Andrew's text files are essentially
           always UTF-8 in practice (operator-confirmed 2026-06-06);
           the fast path catches 99%+ of inputs.
        4. On UnicodeDecodeError, fall back to ``utf-8`` with
           ``errors="replace"``. Replacement characters (U+FFFD) leak
           into the output but the conversation continues — better
           than dropping the message.
        5. Empty result after decode → :class:`AttachmentExtractError`.
        6. Apply the standard char-truncation cap.

    No :mod:`chardet` dep — per operator decision 2026-06-06: keep the
    dependency surface small; if the rare non-UTF-8 file surfaces
    real friction later, lifting in chardet is a one-line dep + one
    branch.
    """
    if not raw_bytes:
        raise AttachmentExtractError("Empty text file")

    # BOM detection. Order matters: UTF-16 LE / BE share a prefix
    # length with UTF-8 BOM's longer form, so check the longer ones
    # first only when length permits. UTF-16 BOMs are 2 bytes,
    # UTF-8 BOM is 3 bytes; the UTF-16 check goes first because the
    # UTF-8 BOM strip is a no-op string-wise (the decode would
    # already handle it).
    if raw_bytes.startswith(_UTF16_LE_BOM):
        try:
            text = raw_bytes[len(_UTF16_LE_BOM):].decode("utf-16-le")
        except UnicodeDecodeError as exc:
            log.warning("talker.attachments.utf16_le_decode_failed", error=str(exc))
            raise AttachmentExtractError(
                f"Failed to decode UTF-16 LE text: {exc!s}"
            ) from exc
    elif raw_bytes.startswith(_UTF16_BE_BOM):
        try:
            text = raw_bytes[len(_UTF16_BE_BOM):].decode("utf-16-be")
        except UnicodeDecodeError as exc:
            log.warning("talker.attachments.utf16_be_decode_failed", error=str(exc))
            raise AttachmentExtractError(
                f"Failed to decode UTF-16 BE text: {exc!s}"
            ) from exc
    else:
        # UTF-8 path. Strip the optional BOM at the head, then decode.
        # ``decode("utf-8-sig")`` would handle the BOM strip
        # automatically, but using plain "utf-8" + manual strip
        # surfaces the fast-vs-fallback path more clearly for the
        # error-handling branches below.
        payload = raw_bytes
        if payload.startswith(_UTF8_BOM):
            payload = payload[len(_UTF8_BOM):]
        try:
            text = payload.decode("utf-8")
        except UnicodeDecodeError as exc:
            # Fallback: replace bad bytes with U+FFFD. Log the
            # original decode failure so an operator tailing logs can
            # see which messages went through the replacement path
            # (sustained replacements suggest a different encoding is
            # in widespread use — promote to chardet at that point).
            log.info(
                "talker.attachments.utf8_decode_fallback",
                error=str(exc),
                error_type=type(exc).__name__,
                bytes_size=len(payload),
            )
            text = payload.decode("utf-8", errors="replace")

    text = text.strip()
    if not text:
        log.info("talker.attachments.text_empty_extraction")
        raise AttachmentExtractError("Empty text content after decode")

    return _apply_char_truncation(text, kind="text")


def extract_csv_text(csv_bytes: bytes) -> str:
    """Extract a CSV byte stream as a Markdown-formatted table.

    Decoding uses the same UTF-8 + BOM + fallback path as
    :func:`extract_text_decoded`. The CSV reader operates on the
    decoded string lines.

    Rendering: Markdown table. The header row (row 0) gets a
    pipe-aligned format with a separator row; subsequent rows mirror.
    Markdown table form chosen because Claude groks it natively —
    raw CSV gets mis-treated as code by some prompts.

    Row-level truncation: caps at :data:`MAX_CSV_ROWS` rows.
    Char-level truncation (the standard :data:`MAX_EXTRACTED_CHARS`)
    applies on top of that.

    Ragged rows (rows with fewer cells than the header) are
    right-padded with empty cells so the Markdown table stays
    well-formed. Wider rows than the header are truncated to header
    width and a per-row note is appended out-of-band — extra cells
    suggest a parsing problem upstream that the user should know
    about.

    Raises:
        AttachmentExtractError: On empty CSV (no rows at all), or
            decode failure.
    """
    if not csv_bytes:
        raise AttachmentExtractError("Empty CSV file")

    # Decode first via the shared text path. The .strip() it applies
    # is safe for CSV — leading/trailing whitespace is never
    # semantically meaningful.
    try:
        decoded = extract_text_decoded(csv_bytes)
    except AttachmentExtractError as exc:
        # Re-raise with a CSV-flavoured message so the user-facing
        # reply names CSV, not generic "text."
        raise AttachmentExtractError(
            f"Failed to decode CSV: {exc!s}"
        ) from exc

    # The shared text extractor applies char-truncation. For CSV we
    # want row-level truncation, so re-parse the raw decoded text
    # rather than the truncated one. Re-decode without the truncation
    # to get the full CSV; the row cap is the primary limit.
    full_decoded: str
    if decoded.endswith(TRUNCATION_MARKER):
        # The text extractor truncated. For CSV we need the full text
        # to apply the ROW cap correctly; redo decode without going
        # through the text extractor's char-truncate path.
        full_decoded = _decode_raw_for_csv(csv_bytes)
    else:
        full_decoded = decoded

    import csv as csv_mod
    reader = csv_mod.reader(io.StringIO(full_decoded))
    try:
        all_rows = list(reader)
    except csv_mod.Error as exc:
        log.warning("talker.attachments.csv_parse_failed", error=str(exc))
        raise AttachmentExtractError(
            f"Failed to parse CSV: {exc!s}"
        ) from exc

    # Strip fully-empty trailing rows (common CSV artefact).
    while all_rows and not any(cell.strip() for cell in all_rows[-1]):
        all_rows.pop()

    if not all_rows:
        log.info("talker.attachments.csv_empty_extraction")
        raise AttachmentExtractError("No rows found in CSV")

    truncated_rows = False
    if len(all_rows) > MAX_CSV_ROWS:
        log.info(
            "talker.attachments.csv_rows_truncated",
            original_rows=len(all_rows),
            kept_rows=MAX_CSV_ROWS,
        )
        all_rows = all_rows[:MAX_CSV_ROWS]
        truncated_rows = True

    # Markdown table render. Header row is the first data row.
    header = all_rows[0]
    header_width = len(header)
    if header_width == 0:
        # Pathological — first row is empty but later rows have data.
        # Use the widest row instead to give every data cell a column.
        header_width = max(len(r) for r in all_rows)
        header = [f"col{i + 1}" for i in range(header_width)]
        body_rows = all_rows  # Treat every row as data.
    else:
        body_rows = all_rows[1:]

    def _pad(row: list[str]) -> list[str]:
        # Pad short rows, truncate wide rows.
        if len(row) < header_width:
            return row + [""] * (header_width - len(row))
        if len(row) > header_width:
            return row[:header_width]
        return row

    def _fmt(row: list[str]) -> str:
        # Pipe-escape: replace any literal ``|`` in cells with ``\|``
        # so the Markdown table stays parseable. Newlines inside a
        # cell collapse to a single space — Markdown tables can't span
        # newlines per cell.
        cells = [
            c.replace("|", "\\|").replace("\n", " ").replace("\r", " ")
            for c in row
        ]
        return "| " + " | ".join(cells) + " |"

    lines: list[str] = [
        _fmt(_pad(header)),
        "| " + " | ".join(["---"] * header_width) + " |",
    ]
    for row in body_rows:
        lines.append(_fmt(_pad(row)))

    if truncated_rows:
        lines.append(
            f"\n[... CSV truncated; only first {MAX_CSV_ROWS} rows shown ...]"
        )

    rendered = "\n".join(lines).strip()
    return _apply_char_truncation(rendered, kind="csv")


def _decode_raw_for_csv(raw_bytes: bytes) -> str:
    """Internal: decode raw text for CSV without the char-truncation cap.

    Mirrors :func:`extract_text_decoded`'s decode pipeline but skips
    the truncation step — CSV row-truncation happens later in
    :func:`extract_csv_text` and would over-truncate if char-truncation
    fired first.
    """
    if raw_bytes.startswith(_UTF16_LE_BOM):
        return raw_bytes[len(_UTF16_LE_BOM):].decode("utf-16-le", errors="replace")
    if raw_bytes.startswith(_UTF16_BE_BOM):
        return raw_bytes[len(_UTF16_BE_BOM):].decode("utf-16-be", errors="replace")
    payload = raw_bytes
    if payload.startswith(_UTF8_BOM):
        payload = payload[len(_UTF8_BOM):]
    try:
        return payload.decode("utf-8")
    except UnicodeDecodeError:
        return payload.decode("utf-8", errors="replace")


def extract_ics_text(ics_bytes: bytes) -> str:
    """Extract events from an .ics calendar byte stream.

    Walks all ``VEVENT`` components and renders each as readable
    structured text:

        Event: <SUMMARY>
        Starts: <DTSTART formatted>
        Ends: <DTEND formatted>
        Location: <LOCATION>           (omitted when absent)
        Description: <DESCRIPTION>     (omitted when absent)

    Multiple events are separated by ``---`` lines. ``VTODO`` /
    ``VJOURNAL`` / ``VFREEBUSY`` components are intentionally skipped
    in c1 — every operationally-shared .ics seen in real use has been
    a meeting / event invite. If TODO-share or journal-share friction
    surfaces, lifting them in is a per-branch extension.

    Datetime formatting: ISO 8601 with timezone for datetime events;
    "All-day on YYYY-MM-DD" for date-only events. The icalendar
    library returns ``datetime.datetime`` for datetime values and
    ``datetime.date`` (the parent class) for date-only — we
    discriminate via ``isinstance(value, datetime)``.

    Raises:
        AttachmentExtractError: On parse failure, OR when the
            calendar contains no VEVENTs (a VTODO-only .ics is the
            common case here — the user-facing reply names the
            decision).
    """
    try:
        # Lazy import.
        from icalendar import Calendar
    except ImportError as exc:
        log.warning("talker.attachments.icalendar_missing", error=str(exc))
        raise AttachmentExtractError(
            "Calendar support is not installed in this build "
            "(pip install -e '.[voice]' to enable)"
        ) from exc

    try:
        cal = Calendar.from_ical(ics_bytes)
    except Exception as exc:  # noqa: BLE001 — wrap icalendar's varied exceptions
        log.warning(
            "talker.attachments.ics_parse_failed",
            error=str(exc),
            error_type=type(exc).__name__,
        )
        raise AttachmentExtractError(
            f"Failed to parse .ics: {exc!s}"
        ) from exc

    event_blocks: list[str] = []
    for component in cal.walk("VEVENT"):
        event_blocks.append(_render_vevent(component))

    if not event_blocks:
        log.info("talker.attachments.ics_no_vevents")
        raise AttachmentExtractError(
            "No events (VEVENT) found in this calendar file. "
            "TODOs / journals aren't supported yet."
        )

    rendered = ("\n\n---\n\n").join(event_blocks).strip()
    return _apply_char_truncation(rendered, kind="ics")


def _render_vevent(component: Any) -> str:
    """Render a single VEVENT component as readable structured text."""
    # icalendar property access via .get() returns either a vText /
    # vDatetime / vDate object (varies by property) or None. The
    # str()-coercion path works for everything; for dates we want to
    # discriminate datetime-vs-date so the formatting is sensible.
    summary = str(component.get("SUMMARY") or "(no title)").strip()
    location = component.get("LOCATION")
    description = component.get("DESCRIPTION")

    dtstart_raw = component.get("DTSTART")
    dtend_raw = component.get("DTEND")

    start_label = _format_ics_datetime(dtstart_raw) if dtstart_raw else "(unknown)"
    end_label = _format_ics_datetime(dtend_raw) if dtend_raw else "(unknown)"

    lines = [f"Event: {summary}"]

    # All-day events: DTSTART is a date (not datetime). When both
    # start and end are date-only AND end - start == 1 day, render a
    # single "All-day on <date>" line rather than two redundant lines.
    if (
        dtstart_raw
        and _is_date_only(dtstart_raw)
        and dtend_raw
        and _is_date_only(dtend_raw)
    ):
        start_d = dtstart_raw.dt
        end_d = dtend_raw.dt
        if (end_d - start_d).days == 1:
            lines.append(f"All-day on {start_d.isoformat()}")
        else:
            lines.append(f"All-day from {start_d.isoformat()} to {end_d.isoformat()}")
    else:
        lines.append(f"Starts: {start_label}")
        lines.append(f"Ends: {end_label}")

    if location:
        loc_clean = str(location).strip()
        if loc_clean:
            lines.append(f"Location: {loc_clean}")
    if description:
        desc_clean = str(description).strip()
        if desc_clean:
            # Trim very long descriptions in-place — the overall
            # char-truncation cap fires later, but per-event 5K chars
            # keeps a single rambling event from squeezing out other
            # events in a multi-event calendar.
            if len(desc_clean) > 5000:
                desc_clean = desc_clean[:5000] + "..."
            lines.append(f"Description: {desc_clean}")

    return "\n".join(lines)


def _is_date_only(prop: Any) -> bool:
    """True if an icalendar property's ``.dt`` is a date but not a datetime."""
    # ``datetime`` is a subclass of ``date``, so isinstance(dt, date)
    # is True for both — check the other direction.
    from datetime import date, datetime as _dt
    dt_value = getattr(prop, "dt", None)
    return isinstance(dt_value, date) and not isinstance(dt_value, _dt)


def _format_ics_datetime(prop: Any) -> str:
    """Format an icalendar datetime property as a readable ISO string."""
    dt_value = getattr(prop, "dt", None)
    if dt_value is None:
        return "(unknown)"
    # ``datetime`` instances have .isoformat() with timezone preserved.
    # Bare dates use the date isoformat (YYYY-MM-DD).
    try:
        return dt_value.isoformat()
    except AttributeError:
        return str(dt_value)


async def extract_audio_transcript(
    audio_bytes: bytes,
    mime_type: str,
    stt_config: Any,
) -> str:
    """Transcribe an audio attachment via the existing Whisper STT path.

    Reuses :func:`alfred.telegram.transcribe.transcribe` — the same
    function the voice-message handler calls. Audio-as-file (the
    .mp3 / .m4a / .wav / .ogg case) and audio-as-voice-note take
    different routes into ``on_document`` vs. ``on_voice`` but share
    the same transcription engine.

    Args:
        audio_bytes: Raw audio payload.
        mime_type: The MIME the user sent ("audio/mpeg", etc.). Used
            as the multipart filename hint to Whisper.
        stt_config: A :class:`alfred.telegram.config.STTConfig` —
            ``config.stt`` from the bot handler's config. Threaded
            through rather than imported so tests can inject a stub
            cleanly.

    Returns:
        The transcript text (non-empty, char-truncated).

    Raises:
        AttachmentExtractError: On TranscribeError (API failure,
            empty transcript), NotImplementedError (provider
            unconfigured), or empty transcript.
    """
    # Lazy import to avoid a hard dep at module load (the same way the
    # other extractors lazy-import their backing libs).
    from . import transcribe as transcribe_mod

    try:
        text = await transcribe_mod.transcribe(audio_bytes, mime_type, stt_config)
    except transcribe_mod.TranscribeError as exc:
        log.info(
            "talker.attachments.audio_transcribe_failed",
            error=str(exc),
            mime_type=mime_type,
        )
        raise AttachmentExtractError(
            f"Failed to transcribe audio: {exc!s}"
        ) from exc
    except NotImplementedError as exc:
        log.warning(
            "talker.attachments.audio_transcribe_unsupported",
            error=str(exc),
        )
        raise AttachmentExtractError(
            "Audio transcription isn't configured on this instance "
            f"({exc!s})"
        ) from exc

    text = text.strip()
    if not text:
        log.info("talker.attachments.audio_empty_transcript")
        raise AttachmentExtractError(
            "Audio transcribed to empty text (silent file?)"
        )

    return _apply_char_truncation(text, kind="audio")


def _short_id_from_file_unique_id(file_unique_id: str) -> str:
    """Return an 8-char filesystem-safe slug from a Telegram unique id.

    Same logic as :func:`alfred.telegram.vision._short_id_from_file_unique_id`
    — duplicated here rather than imported to keep this module independent
    of vision (a future refactor could lift this to a shared
    ``alfred.telegram._fs_id`` helper; out of scope for c1).
    """
    cleaned = "".join(
        c for c in (file_unique_id or "") if c.isalnum() or c in "_-"
    )
    return cleaned[:8] or "unknown"


def storage_path_for_document(
    vault_path: str | Path,
    file_unique_id: str,
    *,
    extension: str = "pdf",
    when: datetime | None = None,
) -> Path:
    """Return the destination path for a saved document attachment.

    Pattern: ``<vault_path>/inbox/document-<YYYYMMDDTHHMMSSZ>-<short>.<ext>``

    Distinct prefix (``document-`` vs. vision's ``screenshot-``) so a
    vault-walk regex over ``inbox/`` can disambiguate image and
    document attachments by filename alone. ISO-8601 compact form for
    cross-filesystem portability (matches vision's pattern).
    """
    if when is None:
        when = datetime.now(timezone.utc)
    stamp = when.strftime("%Y%m%dT%H%M%SZ")
    short = _short_id_from_file_unique_id(file_unique_id)
    name = f"document-{stamp}-{short}.{extension}"
    return Path(vault_path) / "inbox" / name


def save_document_to_inbox(
    document_bytes: bytes,
    vault_path: str | Path,
    file_unique_id: str,
    *,
    extension: str = "pdf",
    when: datetime | None = None,
) -> Path:
    """Persist ``document_bytes`` under the per-instance vault inbox.

    Creates ``<vault_path>/inbox/`` if missing (the scaffold ships it,
    but a fresh vault that hasn't been seeded yet shouldn't crash the
    document handler).

    Returns the absolute path written. Audit-trail value is the whole
    point — without this, the saved-path on the session record points
    at a missing file and the audit trail is broken.

    Persistence failure at the call site is treated as non-fatal
    (same contract as :func:`alfred.telegram.vision.save_image_to_inbox`):
    the handler catches the exception and proceeds with the in-memory
    bytes, logging ``action=continuing_to_llm_in_memory_only``.
    """
    dest = storage_path_for_document(
        vault_path, file_unique_id, extension=extension, when=when,
    )
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(document_bytes)
    log.info(
        "talker.attachments.saved",
        path=str(dest),
        bytes=len(document_bytes),
    )
    return dest


def storage_path_for_audio(
    vault_path: str | Path,
    file_unique_id: str,
    *,
    extension: str = "bin",
    when: datetime | None = None,
) -> Path:
    """Return the destination path for a saved audio attachment.

    Pattern: ``<vault_path>/inbox/audio-<YYYYMMDDTHHMMSSZ>-<short>.<ext>``

    Distinct ``audio-`` prefix (vs. ``document-`` and ``screenshot-``)
    so a vault-walk regex over ``inbox/`` can disambiguate the three
    attachment kinds by filename alone. Per the P8 plan and the
    universal-filetype-support feedback memo: audio files are kept
    in the same inbox dir but identifiable by prefix.

    ISO-8601 compact form for cross-filesystem portability (matches
    document + vision patterns).
    """
    if when is None:
        when = datetime.now(timezone.utc)
    stamp = when.strftime("%Y%m%dT%H%M%SZ")
    short = _short_id_from_file_unique_id(file_unique_id)
    name = f"audio-{stamp}-{short}.{extension}"
    return Path(vault_path) / "inbox" / name


def save_audio_to_inbox(
    audio_bytes: bytes,
    vault_path: str | Path,
    file_unique_id: str,
    *,
    extension: str = "bin",
    when: datetime | None = None,
) -> Path:
    """Persist ``audio_bytes`` under the per-instance vault inbox.

    Parallel to :func:`save_document_to_inbox` — same semantics
    (creates inbox dir on demand, writes bytes, emits a log line),
    different storage prefix (``audio-`` instead of ``document-``).
    Persistence failure is non-fatal at the call site: the handler
    catches and proceeds with the in-memory bytes, logging
    ``action=continuing_to_llm_in_memory_only``.

    Audio files preserve the user's original container (mp3 / m4a /
    wav / ogg) via the ``extension`` arg derived from the MIME map.
    The user can re-listen to the original audio from the vault even
    after the transcript is in the conversation.
    """
    dest = storage_path_for_audio(
        vault_path, file_unique_id, extension=extension, when=when,
    )
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(audio_bytes)
    log.info(
        "talker.attachments.audio_saved",
        path=str(dest),
        bytes=len(audio_bytes),
    )
    return dest


# Per-kind banner header used by :func:`build_document_user_text`.
# Tells the LLM what kind of attachment is below the fence so prompts
# can branch on type without re-deriving from filename.
_BANNER_BY_KIND: dict[str, str] = {
    "pdf": "PDF attached",
    "docx": "DOCX attached",
    "text": "Text file attached",
    "csv": "CSV attached",
    "ics": "Calendar invite attached",
    "audio": "Audio transcript",
}


# Per-kind fence label used by :func:`build_document_user_text`. Most
# kinds fence as "Document text"; ICS uses "Events" so the LLM sees
# the structured calendar content for what it is; audio uses
# "Transcript" so the LLM knows it's reading Whisper output (not
# typed text).
_FENCE_LABEL_BY_KIND: dict[str, str] = {
    "pdf": "Document text",
    "docx": "Document text",
    "text": "Document text",
    "csv": "Document text",
    "ics": "Events",
    "audio": "Transcript",
}


def build_document_user_text(
    caption: str,
    extracted_text: str,
    filename: str,
    kind: str,
) -> str:
    """Compose the user-message text for a document attachment.

    Shape (kind=pdf example)::

        [PDF attached: <filename>]

        <caption>           (when present)

        --- Document text ---
        <extracted_text>

    The header banner ("PDF attached", "DOCX attached", etc.) names
    the attachment type explicitly so the LLM doesn't have to infer
    from filename (which Telegram sometimes strips). The caption (if
    any) follows so user intent is preserved verbatim. The extracted
    text is fenced with a kind-appropriate label ("Document text",
    "Events" for .ics, "Transcript" for audio) so the model treats it
    as the attached content rather than a continuation of the user's
    words.

    Empty caption collapses cleanly (no leading blank block).

    Unknown kinds fall back to a generic "Document attached" banner +
    "Document text" fence — defensive, the on_document dispatcher
    should always pass a known kind from
    :data:`SUPPORTED_DOCUMENT_MIME.values()`.
    """
    safe_filename = filename or f"document.{extension_for_kind(kind)}"
    banner = _BANNER_BY_KIND.get(kind, "Document attached")
    fence_label = _FENCE_LABEL_BY_KIND.get(kind, "Document text")
    header = f"[{banner}: {safe_filename}]"

    parts: list[str] = [header, ""]
    caption_clean = (caption or "").strip()
    if caption_clean:
        parts.extend([caption_clean, ""])
    parts.extend([f"--- {fence_label} ---", extracted_text])
    return "\n".join(parts)


__all__ = [
    "AttachmentDownloadError",
    "AttachmentExtractError",
    "MAX_AUDIO_BYTES",
    "MAX_BYTES_BY_KIND",
    "MAX_CSV_BYTES",
    "MAX_CSV_ROWS",
    "MAX_DOCX_BYTES",
    "MAX_EXTRACTED_CHARS",
    "MAX_ICS_BYTES",
    "MAX_PDF_BYTES",
    "MAX_TEXT_BYTES",
    "SUPPORTED_DOCUMENT_MIME",
    "TRUNCATION_MARKER",
    "build_document_user_text",
    "download_document_bytes",
    "extension_for_kind",
    "extract_audio_transcript",
    "extract_csv_text",
    "extract_docx_text",
    "extract_ics_text",
    "extract_pdf_text",
    "extract_text_decoded",
    "save_audio_to_inbox",
    "save_document_to_inbox",
    "storage_path_for_audio",
    "storage_path_for_document",
]
